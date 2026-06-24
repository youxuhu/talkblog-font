from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 页面设置（匹配原文档） ──
section = doc.sections[0]
section.page_width = Emu(7560310)
section.page_height = Emu(10692130)
section.top_margin = Emu(914400)
section.bottom_margin = Emu(914400)
section.left_margin = Emu(1143000)
section.right_margin = Emu(1143000)

# ── 默认样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_run_font(run, size=Pt(14), bold=False, font_name='宋体'):
    run.font.size = size
    run.font.name = font_name
    run.bold = bold
    rpr = run.element.rPr
    if rpr is None:
        rpr = run.element.makeelement(qn('w:rPr'), {})
        run.element.insert(0, rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def add_normal_para(text, bold=False, indent_left=False, align=None, font_name='宋体', size=Pt(14)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, font_name=font_name)
    if indent_left:
        p.paragraph_format.first_line_indent = Emu(355600)
    if align:
        p.alignment = align
    return p

# ════════════════════════════════
# 标题（标题行跟在原文档中没有，用大号字加粗居中）
# ════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('第八组第一次讨论纪要')
set_run_font(run, size=Pt(18), bold=True)
title_p.paragraph_format.space_after = Pt(6)

# ════════════════════════════════
# 信息行
# ════════════════════════════════
add_normal_para('时间：2026/4/16')
add_normal_para('成员：')
add_normal_para('1.胡有绪（组长）  2024081105', indent_left=True)
add_normal_para('2.陈锐  2024081108', indent_left=True)
add_normal_para('3.陈朗  2024081106', indent_left=True)
add_normal_para('4.彭晋弘  2024081109', indent_left=True)
add_normal_para('5.杨家豪  2024081110', indent_left=True)
add_normal_para('地点：双中心B511')

# ════════════════════════════════
# 讨论主题
# ════════════════════════════════
add_normal_para('讨论主题：项目的基本框架', bold=True)

# ════════════════════════════════
# 讨论结果
# ════════════════════════════════
add_normal_para('讨论结果：', bold=True)

# ── 一、技术栈选型 ──
add_normal_para('一、技术栈选型', bold=True)
table = doc.add_table(rows=7, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['层面', '选型', '说明']
data = [
    ['前端框架', 'Vue 3 + Composition API', '逻辑复用优于 Options API，适合组件化开发'],
    ['构建工具', 'Vite 8', '快速 HMR，现代化构建'],
    ['状态管理', 'Pinia', '模块化 store，TS 支持好，替代 Vuex'],
    ['路由', 'Vue Router 5', 'SPA 路由管理，支持嵌套路由与导航守卫'],
    ['UI 框架', 'Element Plus + Pixel UI', '双框架共存，过渡阶段'],
    ['后端代理', 'Vite proxy → localhost:8080', '开发环境代理 API 请求'],
]
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_run_font(run, size=Pt(14), bold=True)
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        set_run_font(run, size=Pt(14))

# ── 二、项目目录结构 ──
add_normal_para('二、项目目录结构', bold=True)
code_lines = [
    'talkblog-font/',
    '├── src/',
    '│   ├── main.js              # 应用入口，注册插件',
    '│   ├── App.vue              # 根组件，导航栏 + 路由视图',
    '│   ├── router/index.js      # 路由配置 + 守卫',
    '│   ├── stores/              # Pinia 状态管理（7 个 store）',
    '│   ├── services/            # API 服务层（8 个 service）',
    '│   ├── views/               # 页面组件（16 个视图）',
    '│   ├── components/          # 公共组件（4 个）',
    '│   ├── config/              # 配置文件',
    '│   └── styles/              # 全局样式（主题切换）',
    '├── doc/                     # 项目文档',
    '└── public/                  # 静态资源',
]
for line in code_lines:
    add_normal_para(line, font_name='Courier New', size=Pt(12))

# ── 三、功能模块划分 ──
add_normal_para('三、功能模块划分', bold=True)
modules = [
    '1. 认证模块 — 人脸/密码/邮箱登录，JWT 鉴权，路由守卫',
    '2. 博客模块 — 博客 CRUD、版本管理、列表/详情/编辑器',
    '3. 评论社交模块 — 评论回复/点赞/编辑/举报，支持 emoji、@提及、图片',
    '4. 专栏模块 — 专栏 CRUD、列表/详情/编辑器',
    '5. 通知模块 — WebSocket 实时推送，通知列表',
    '6. 后台管理 — 用户管理、评论审核、统计、举报处理',
    '7. 其他 — 书签、关注、主题切换、分类配置',
]
for m in modules:
    add_normal_para(m, indent_left=True)

# ── 四、数据流架构 ──
add_normal_para('四、数据流架构', bold=True)
flow_lines = [
    'View (Views/Components)',
    '  → Store (Pinia)',
    '    → Service (fetch)',
    '      → Backend API (/api/*)',
]
for line in flow_lines:
    add_normal_para(line, font_name='Courier New', size=Pt(12))

items = [
    '- View 层负责 UI 展示与交互',
    '- Store 层管理全局状态',
    '- Service 层封装 HTTP 请求与业务逻辑',
    '- 后端通过 Vite proxy 代理到 localhost:8080',
]
for item in items:
    add_normal_para(item)

# ── 五、路由守卫策略 ──
add_normal_para('五、路由守卫策略', bold=True)
add_normal_para('三层校验机制：')
guards = [
    '- requiresAuth → 未登录跳转首页',
    '- requiresAdmin → 非管理员提示权限不足',
    '- 已登录访问首页 → 自动跳转 /welcome',
]
for g in guards:
    add_normal_para(g)

doc.save('doc/第八组讨论纪要.docx')
print('OK')
