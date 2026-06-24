from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Emu(7560310)
section.page_height = Emu(10692130)
section.top_margin = Emu(914400)
section.bottom_margin = Emu(914400)
section.left_margin = Emu(1143000)
section.right_margin = Emu(1143000)

# ── 默认样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_run_font(run, size=Pt(14), bold=False, font_name='宋体'):
    run.font.size = size
    run.font.name = font_name
    run.bold = bold
    rpr = run.element.rPr
    if rpr is None:
        rpr = run.element.makeelement(qn('w:rPr'), {})
        run.element.insert(0, rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def add_normal_para(text, bold=False, indent_left=False, align=None, font_name='宋体', size=Pt(14)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, font_name=font_name)
    if indent_left:
        p.paragraph_format.first_line_indent = Emu(355600)
    if align:
        p.alignment = align
    return p

def add_table_with_data(headers, data, col_widths=None):
    table = doc.add_table(rows=1 + len(data), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=Pt(14), bold=True)
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, size=Pt(14))
    return table

# ═══════════ 标题 ═══════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('第八组第二次讨论纪要')
set_run_font(run, size=Pt(18), bold=True)
title_p.paragraph_format.space_after = Pt(6)

# ═══════════ 信息 ═══════════
add_normal_para('时间：2026/4/23')
add_normal_para('成员：')
add_normal_para('1.胡有绪（组长）  2024081105', indent_left=True)
add_normal_para('2.陈锐  2024081108', indent_left=True)
add_normal_para('3.陈朗  2024081106', indent_left=True)
add_normal_para('4.彭晋弘  2024081109', indent_left=True)
add_normal_para('5.杨家豪  2024081110', indent_left=True)
add_normal_para('地点：双中心B511')

add_normal_para('讨论主题：项目功能结构与实现方案', bold=True)
add_normal_para('讨论结果：', bold=True)

# ── 一、认证模块实现 ──
add_normal_para('一、认证模块实现', bold=True)
add_table_with_data(
    ['功能点', '实现方式', '关键技术'],
    [
        ['人脸注册', '摄像头采集 → base64 → POST /api/auth/register', 'WebRTC getUserMedia'],
        ['人脸登录', '摄像头采集 → base64 → POST /api/auth/login', 'Face API 后端比对'],
        ['密码登录', '表单提交 → POST /api/auth/password-login', 'JWT 令牌返回'],
        ['会话管理', 'sessionStorage 存储 authState', 'JWT 解码 + 路由守卫'],

        ['权限校验', 'isAuthenticated() / isAdminUser()', '三级路由守卫'],
    ]
)

add_normal_para('认证流程说明：')
add_normal_para('用户在 HomeView 选择登录模式 → 调用对应 Service → 后端返回 JWT → 存入 sessionStorage（关闭标签页自动过期） → Pinia auth store 同步状态 → 路由守卫根据 meta 标签（requiresAuth / requiresAdmin）控制页面访问。', indent_left=True)

# ── 二、状态管理层（Pinia Store） ──
add_normal_para('二、状态管理层（Pinia Store）', bold=True)
add_table_with_data(
    ['Store', '核心状态', '主要 Action 数', '职责'],
    [
        ['auth', 'authState / loading', '5', '认证状态管理'],
        ['blog', 'blogs / currentBlog / total', '9', '博客 CRUD 与列表缓存'],
        ['comment', 'comments / adminComments / reports / stats', '16', '评论全生命周期'],
        ['series', 'seriesList / currentSeries', '8', '专栏管理'],
        ['notification', 'notifications / unreadCount', '7', '通知列表与 WebSocket'],
        ['theme', 'isDark', '2', '暗色模式切换'],
        ['counter', 'count', '1', '示例/遗留 store'],
    ]
)

add_normal_para('设计特点：每个 store 独立职责，通过解构 service 层 API 实现数据获取，统一使用 ref/reactive 响应式状态，getter 通过 computed 派生。', indent_left=True)

# ── 三、服务层（Service） ──
add_normal_para('三、服务层（Service）', bold=True)
add_table_with_data(
    ['Service 文件', 'API 端点数量', '核心功能'],
    [
        ['auth.js', '9 个', '认证 + 用户管理 + session 工具'],
        ['blog.js', '7 个', '博客 CRUD + 分类查询'],
        ['comment.js', '16 个', '评论增删改查 + 管理审核 + 举报 + 统计'],
        ['series.js', '6 个', '专栏管理 + 专栏内博客查询'],
        ['notification.js', '4 + WebSocket', '通知 CRUD + 实时推送'],
        ['bookmark.js', '3 个', '书签切换与查询'],
        ['follow.js', '5 个', '关注/取关 + 粉丝/关注列表'],
        ['blogVersion.js', '1 个', '博客版本历史'],
    ]
)

add_normal_para('公共模式分析：')
add_normal_para('1. 统一使用 buildUrl() 构造带查询参数的 URL', indent_left=True)
add_normal_para('2. 统一使用 requestJson() 封装 fetch，自动附加 Authorization 头', indent_left=True)
add_normal_para('3. 统一 JSON 解析与错误处理（throw on !response.ok 或 success===false）', indent_left=True)
add_normal_para('4. 存在问题：8 个 service 存在大量重复代码，需抽取公共 HTTP 模块', indent_left=True)

# ── 四、评论系统架构 ──
add_normal_para('四、评论系统架构', bold=True)
add_normal_para('1. 评论模型：支持多级嵌套回复（递归组件，最大深度 4 层）', indent_left=True)
add_normal_para('2. 交互功能：评论/回复/编辑/删除/点赞/举报/置顶/图片上传/emoji/@提及', indent_left=True)
add_normal_para('3. 审核机制：用户评论 → 可选审核 → 管理员审核/批量操作', indent_left=True)
add_normal_para('4. 管理后台：评论列表/搜索/状态筛选/统计报表/举报处理', indent_left=True)
add_normal_para('5. 当前模式：写后读（操作后 emit refresh 全量重载），建议改为乐观更新', indent_left=True)

add_normal_para('CommentItem 组件功能：')
add_normal_para('• 确定性渐变头像（用户名 hash 映射 8 色）', indent_left=True)
add_normal_para('• @提及高亮 + 可点击跳转', indent_left=True)
add_normal_para('• 相对时间显示（刚刚/X分钟前/X小时前/...）', indent_left=True)
add_normal_para('• 图片网格预览 + 点击放大', indent_left=True)
add_normal_para('• 已编辑/已置顶/待审核 状态标记', indent_left=True)

# ── 五、通知系统设计 ──
add_normal_para('五、通知系统设计', bold=True)
add_normal_para('1. 推送方式：原生 WebSocket（ws[s]://服务器/ws/notifications?token=xxx）', indent_left=True)
add_normal_para('2. 自动重连：断开 5s 重试，错误 10s 重试', indent_left=True)
add_normal_para('3. 通知类型：comment/reply/like/mention/system/blog_published/scheduled_publish', indent_left=True)
add_normal_para('4. UI 组件：NotificationBell 顶部铃铛 → 下拉列表 → 标记已读/全部已读 → 跳转详情', indent_left=True)
add_normal_para('5. 未读计数：服务端返回 + 本地实时递减', indent_left=True)

# ── 六、UI 与样式 ──
add_normal_para('六、UI 与样式', bold=True)
add_normal_para('1. UI 框架：Element Plus（主要组件库）+ Pixel UI（@mmt817/pixel-ui，过渡产物）', indent_left=True)
add_normal_para('2. 主题系统：CSS 变量驱动，:root（亮色）+ .dark（暗色）两套变量覆盖', indent_left=True)
add_normal_para('3. 暗色模式：覆盖 Element Plus 全部组件样式（表格/分页/输入框/选择器/标签/弹窗等）', indent_left=True)
add_normal_para('4. 色彩系统：紫色主色调 #7c4dff，绿色成功 #2f855a，橙色警告 #ed8936，红色危险 #e53e3e', indent_left=True)
add_normal_para('5. 问题：双 UI 框架共存增加约 30% 样式覆盖代码，应考虑统一', indent_left=True)

# ── 七、路由与页面结构 ──
add_normal_para('七、路由与页面结构', bold=True)
add_table_with_data(
    ['页面', '路由路径', '权限要求', '功能'],
    [
        ['首页', '/', '无', '登录/注册选择'],
        ['欢迎页', '/welcome', '需登录', '用户首页'],
        ['博客列表', '/blogs', '无', '博客浏览'],
        ['博客详情', '/blog/:id', '无', '阅读+评论'],
        ['博客编辑', '/blog/editor/:id?', '需登录', '编写/修改'],
        ['专栏列表', '/series', '无', '专栏浏览'],
        ['专栏详情', '/series/:id', '无', '专栏内容'],
        ['通知', '/notifications', '需登录', '通知列表'],
        ['书签', '/bookmarks', '需登录', '收藏博客'],
        ['个人资料', '/profile', '需登录', '个人信息维护'],
        ['管理后台', '/admin/*', '需管理员', '用户/评论/统计/举报'],
    ]
)

doc.save('doc/第八组第二次讨论纪要.docx')
print('OK')
